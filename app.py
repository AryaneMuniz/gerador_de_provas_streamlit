import streamlit as st
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import date
from io import BytesIO

st.set_page_config(page_title="Gerador de Provas", layout="centered")
st.title("📝 Gerador de Provas Escolares")

if "questoes" not in st.session_state:
    st.session_state.questoes = []
if "editando_index" not in st.session_state:
    st.session_state.editando_index = None
if "texto_questao" not in st.session_state:
    st.session_state.texto_questao = ""
if "imagem_questao" not in st.session_state:
    st.session_state.imagem_questao = None
if "opcoes" not in st.session_state:
    st.session_state.opcoes = {"A": "", "B": "", "C": "", "D": ""}
if "tipo_questao" not in st.session_state:
    st.session_state.tipo_questao = "Dissertativa"

st.sidebar.header("Configurações do Cabeçalho")
logo_escola = st.sidebar.file_uploader("📌 Logo da Escola (PNG/JPG)", type=["png", "jpg", "jpeg"])

with st.form("dados_prova"):
    st.subheader("📋 Dados da Prova")
    nome_professor = st.text_input("Nome do Professor*", placeholder="Obrigatório")
    disciplina = st.text_input("Disciplina*", placeholder="Obrigatório")
    serie = st.selectbox("Série/Turma*", [
        "1º ano - Ensino Fundamental", "2º ano - Ensino Fundamental", 
        "3º ano - Ensino Fundamental", "4º ano - Ensino Fundamental",
        "5º ano - Ensino Fundamental", "6º ano - Ensino Fundamental",
        "7º ano - Ensino Fundamental", "8º ano - Ensino Fundamental",
        "9º ano - Ensino Fundamental", "1º ano - Ensino Médio",
        "2º ano - Ensino Médio", "3º ano - Ensino Médio"
    ])
    bimestre = st.selectbox("Bimestre*", ["1º Bimestre", "2º Bimestre", "3º Bimestre", "4º Bimestre"])
    data_prova = st.date_input("Data da Prova*", value=date.today())
    st.form_submit_button("Salvar Configurações")

st.subheader("✍️ Editor de Questões")
st.session_state.tipo_questao = st.radio("Tipo de Questão*", ["Dissertativa", "Múltipla Escolha"], horizontal=True)

texto_questao = st.text_area("Texto da Questão*", height=150, value=st.session_state.texto_questao, placeholder="Digite o enunciado da questão...")
imagem_questao = st.file_uploader("Imagem de Apoio (opcional)", type=["png", "jpg", "jpeg"])

if st.session_state.tipo_questao == "Múltipla Escolha":
    st.markdown("**Opções de Resposta:**")
    col1, col2 = st.columns(2)
    with col1:
        st.session_state.opcoes["A"] = st.text_input("Opção A*", value=st.session_state.opcoes["A"], placeholder="Texto da opção A")
        st.session_state.opcoes["C"] = st.text_input("Opção C*", value=st.session_state.opcoes["C"], placeholder="Texto da opção C")
    with col2:
        st.session_state.opcoes["B"] = st.text_input("Opção B*", value=st.session_state.opcoes["B"], placeholder="Texto da opção B")
        st.session_state.opcoes["D"] = st.text_input("Opção D*", value=st.session_state.opcoes["D"], placeholder="Texto da opção D")

col_salvar, col_limpar = st.columns(2)
with col_salvar:
    if st.button("💾 Salvar Questão", use_container_width=True):
        if not texto_questao.strip():
            st.error("O texto da questão é obrigatório!")
        elif st.session_state.tipo_questao == "Múltipla Escolha" and any(not opcao.strip() for opcao in st.session_state.opcoes.values()):
            st.error("Preencha todas as opções de múltipla escolha!")
        else:
            nova_questao = {
                "texto": texto_questao,
                "tipo": st.session_state.tipo_questao,
                "imagem": imagem_questao.read() if imagem_questao else None,
                "opcoes": st.session_state.opcoes.copy() if st.session_state.tipo_questao == "Múltipla Escolha" else None
            }
            
            if st.session_state.editando_index is not None:
                st.session_state.questoes[st.session_state.editando_index] = nova_questao
                st.success("Questão atualizada com sucesso!")
                st.session_state.editando_index = None
            else:
                st.session_state.questoes.append(nova_questao)
                st.success("Questão adicionada com sucesso!")
            
            st.session_state.texto_questao = ""
            st.session_state.opcoes = {"A": "", "B": "", "C": "", "D": ""}
            st.rerun()

with col_limpar:
    if st.button("♻️ Limpar Campos", use_container_width=True):
        st.session_state.texto_questao = ""
        st.session_state.opcoes = {"A": "", "B": "", "C": "", "D": ""}
        st.session_state.editando_index = None
        st.rerun()

st.subheader("📚 Questões Adicionadas")
st.caption(f"Total: {len(st.session_state.questoes)} questões")

if not st.session_state.questoes:
    st.info("Nenhuma questão adicionada ainda. Use o editor acima para começar.")
else:
    for idx, questao in enumerate(st.session_state.questoes):
        with st.expander(f"Questão {idx + 1}: {questao['texto'][:50]}...", expanded=False):
            st.markdown(f"**Enunciado:** {questao['texto']}")
            
            if questao["imagem"]:
                st.image(BytesIO(questao["imagem"]), width=300)
            
            if questao["tipo"] == "Múltipla Escolha":
                st.markdown("**Alternativas:**")
                cols = st.columns(2)
                with cols[0]:
                    st.markdown(f"**A)** {questao['opcoes']['A']}")
                    st.markdown(f"**C)** {questao['opcoes']['C']}")
                with cols[1]:
                    st.markdown(f"**B)** {questao['opcoes']['B']}")
                    st.markdown(f"**D)** {questao['opcoes']['D']}")
            else:
                st.markdown("**Tipo:** Dissertativa")
                st.markdown("_" * 50)
            
            if st.button("✏️ Editar", key=f"edit_{idx}"):
                st.session_state.editando_index = idx
                st.session_state.texto_questao = questao["texto"]
                st.session_state.tipo_questao = questao["tipo"]
                if questao["opcoes"]:
                    st.session_state.opcoes = questao["opcoes"].copy()
                st.rerun()
            
            if st.button("🗑️ Excluir", key=f"del_{idx}"):
                st.session_state.questoes.pop(idx)
                st.success("Questão removida!")
                st.rerun()

st.subheader("📤 Exportar Prova")
if st.button("💾 Gerar Documento Word", use_container_width=True):
    if not st.session_state.questoes:
        st.error("Adicione pelo menos uma questão antes de exportar!")
    elif not all([nome_professor, disciplina, serie, bimestre]):
        st.error("Preencha todos os campos obrigatórios nos dados da prova!")
    else:
        try:
            doc = Document()
            style = doc.styles['Normal']
            style.font.name = 'Arial'
            style.font.size = Pt(12)
            
            if logo_escola:
                logo_escola.seek(0)
                doc.add_picture(logo_escola, width=Inches(1.5))
                last_paragraph = doc.paragraphs[-1] 
                last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            titulo = doc.add_paragraph(f"PROVA DE {disciplina.upper()} - {bimestre.upper()}")
            titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
            titulo.runs[0].bold = True
            
            doc.add_paragraph(f"Professor: {nome_professor}")
            doc.add_paragraph(f"Turma: {serie}")
            doc.add_paragraph(f"Data: {data_prova.strftime('%d/%m/%Y')}")
            doc.add_paragraph("\n")
            
            for idx, questao in enumerate(st.session_state.questoes, 1):
                para = doc.add_paragraph()
                para.add_run(f"{idx}. ").bold = True
                para.add_run(questao["texto"])
                
                if questao["imagem"]:
                    try:
                        doc.add_picture(BytesIO(questao["imagem"]), width=Inches(4.5))
                        doc.paragraphs[-1].alignment = WD_ALIGN_PARAGRAPH.CENTER
                    except:
                        doc.add_paragraph("[Erro ao carregar imagem]")
                
                if questao["tipo"] == "Múltipla Escolha":
                    for letra, texto in questao["opcoes"].items():
                        doc.add_paragraph(f"{letra}) {texto}")
                else:
                    for _ in range(3):
                        doc.add_paragraph("_" * 60)
                
                doc.add_paragraph()
            
            buffer = BytesIO()
            doc.save(buffer)
            buffer.seek(0)
            
            nome_arquivo = f"Prova_{disciplina}_{serie}_{bimestre}.docx".replace(" ", "_")
            st.download_button(
                "⬇️ Baixar Prova em Word",
                data=buffer,
                file_name=nome_arquivo,
                mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                use_container_width=True
            )
        except Exception as e:
            st.error(f"Erro ao gerar documento: {str(e)}")
